# -*- coding: utf-8 -*-
"""
laudo_docx.py — Geração do laudo em Word (.docx) via python-docx.

Design system alinhado ao laudo_detalhado.html:
  Títulos    : Times New Roman (aproxima Playfair Display), navy #12345e
  Corpo      : Calibri 11pt (aproxima Source Sans 3)
  Tabelas    : cabeçalho navy #12345e, alternância #EFF4F9
  Alertas    : caixa com borda esquerda colorida (ok/ambar/critico)
  Capa       : página exclusiva com título, espólio e responsável

Sistema OrgAudi · ORGATEC CONTABILIDADE E AUDITORIA
"""
from __future__ import annotations

from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .dominio import Espolio, Regime, StatusApuracao
from .motor import MotorApuracao


# ───────────────────────────────────────────────────────────────────────
# Paleta
# ───────────────────────────────────────────────────────────────────────

_NAVY   = RGBColor(0x12, 0x34, 0x5e)
_BLUE   = RGBColor(0x1f, 0x7f, 0xb8)
_CYAN   = RGBColor(0x38, 0xc4, 0xe6)
_OK     = RGBColor(0x2f, 0x7d, 0x4f)
_WARN   = RGBColor(0xb7, 0x79, 0x1f)
_ERR    = RGBColor(0xb3, 0x3a, 0x3a)
_GRAY   = RGBColor(0x51, 0x61, 0x6f)
_WHITE  = RGBColor(0xff, 0xff, 0xff)
_BG_TBL = "EFF4F9"
_BG_KEY = "EEF3F8"
_BG_TOT = "DCE9F4"
_BG_HDR = "12345E"
_BG_OK  = "E3F1E8"
_BG_AMB = "FBF1DA"
_BG_ERR = "FBF0EF"


# ───────────────────────────────────────────────────────────────────────
# Helpers de XML / estilo
# ───────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, side: str, color: str, sz: int = 12):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:color"), color)
    tcBorders.append(border)


def _add_left_bar(paragraph, hex_color: str, width_pt: int = 12):
    """Adiciona borda esquerda colorida num parágrafo (simula .aviso)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))
    left.set(qn("w:color"), hex_color)
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)


def _set_para_shading(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


# ───────────────────────────────────────────────────────────────────────
# Construtores de elementos
# ───────────────────────────────────────────────────────────────────────

def _m(v) -> str:
    if v is None:
        return "—"
    s = f"{v:,.2f}"
    return "R$ " + s.replace(",", "X").replace(".", ",").replace("X", ".")


def _d(d) -> str:
    if d is None:
        return "—"
    return d.strftime("%d/%m/%Y") if hasattr(d, "strftime") else str(d)


def _run(paragraph, text: str, bold=False, italic=False,
         color: RGBColor | None = None, size_pt: int = 11,
         font: str = "Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = _NAVY
    if level == 1:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
        # borda inferior
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:color"), "12345E")
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        run.font.size = Pt(11.5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p


def _aviso(doc: Document, tag: str, texto: str,
           tipo: str = "info"):
    color_map = {"ok": "2F7D4F", "ambar": "B7791F",
                 "critico": "B33A3A", "info": "1F7FB8"}
    bg_map    = {"ok": _BG_OK, "ambar": _BG_AMB,
                 "critico": _BG_ERR, "info": "EEF6FB"}
    cor = color_map.get(tipo, "1F7FB8")
    bg  = bg_map.get(tipo, "EEF6FB")

    p_tag = doc.add_paragraph()
    _set_para_shading(p_tag, bg)
    _add_left_bar(p_tag, cor, 18)
    p_tag.paragraph_format.left_indent  = Cm(0.4)
    p_tag.paragraph_format.space_after  = Pt(0)
    r = p_tag.add_run(tag.upper() + "  ")
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(cor)

    p_body = doc.add_paragraph()
    _set_para_shading(p_body, bg)
    _add_left_bar(p_body, cor, 18)
    p_body.paragraph_format.left_indent = Cm(0.4)
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after  = Pt(8)
    _run(p_body, texto, size_pt=10)


def _table_header_row(row, cols: list[str]):
    for i, col in enumerate(cols):
        cell = row.cells[i]
        _set_cell_bg(cell, _BG_HDR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(col)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = _WHITE


def _table_data_row(row, values: list[str | None],
                    aligns: list[str] | None = None,
                    bg: str | None = None, bold: bool = False,
                    color: RGBColor | None = None):
    aligns = aligns or ["left"] * len(values)
    al_map = {"left": WD_ALIGN_PARAGRAPH.LEFT,
              "right": WD_ALIGN_PARAGRAPH.RIGHT,
              "center": WD_ALIGN_PARAGRAPH.CENTER}
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = al_map.get(aligns[i], WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(val or "—")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.bold = bold
        if color:
            run.font.color.rgb = color


# ───────────────────────────────────────────────────────────────────────
# Construção das seções
# ───────────────────────────────────────────────────────────────────────

def _bloco_capa(doc: Document, espolio: Espolio):
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("LAUDO DE APURAÇÃO DE IR\nSOBRE GANHO DE CAPITAL — ESPÓLIO")
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.font.color.rgb = _NAVY
    run.bold = True

    doc.add_paragraph()  # espaço

    # Linha decorativa
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("─" * 40)
    run2.font.color.rgb = _BLUE

    doc.add_paragraph()

    # Espólio
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, f"{espolio.nome}\n", bold=True, color=_NAVY, size_pt=14,
         font="Times New Roman")
    _run(p3, f"CPF {espolio.cpf_falecido}", color=_GRAY, size_pt=11)

    doc.add_paragraph()

    # Meta
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p4, f"Óbito: {_d(espolio.data_obito)}  ·  Partilha: {_d(espolio.data_partilha)}\n",
         color=_GRAY, size_pt=10)
    _run(p4, f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n",
         color=_GRAY, size_pt=10)
    _run(p4, "ORGATEC Contabilidade e Auditoria", color=_BLUE, size_pt=10)

    doc.add_page_break()


def _bloco_resumo(doc: Document, resumo: dict, espolio: Espolio):
    _heading(doc, "1. Resumo Executivo")

    # Métricas rápidas
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"

    hdrs = ["IR Total do Espólio", "Bens Conclusivos", "Bens Pendentes", "Status"]
    _table_header_row(tbl.rows[0], hdrs)

    ir_total = resumo.get("ir_espolio_total", 0.0)
    conclusivo = resumo.get("conclusivo", False)
    vals = [
        _m(ir_total),
        str(resumo.get("bens_conclusivos", 0)),
        str(resumo.get("bens_pendentes", 0)),
        "CONCLUSIVO" if conclusivo else "PENDENTE",
    ]
    _table_data_row(tbl.rows[1], vals, bold=True, color=_NAVY,
                    bg=_BG_KEY)

    doc.add_paragraph()

    if not conclusivo:
        _aviso(doc, "Atenção",
               f"{resumo.get('bens_pendentes',0)} bem(ns) com pendência. "
               "O IR total NÃO é definitivo enquanto houver lacunas.", "ambar")

    # Participantes
    _heading(doc, "Participantes", level=2)
    tbl2 = doc.add_table(rows=1 + len(espolio.herdeiros), cols=4)
    tbl2.style = "Table Grid"
    _table_header_row(tbl2.rows[0], ["Nome", "CPF", "Fração", "Tipo"])
    for i, h in enumerate(espolio.herdeiros, 1):
        bg = _BG_TBL if i % 2 == 0 else None
        _table_data_row(tbl2.rows[i], [
            h.nome, h.cpf or "—",
            f"{h.fracao_monte*100:.2f}%",
            "Meeiro" if h.eh_meeiro else "Herdeiro",
        ], bg=bg)

    doc.add_paragraph()


def _bloco_apuracao(doc: Document, espolio: Espolio):
    _heading(doc, "2. Apuração por Bem")

    cols = ["Bem", "Regime", "Status", "Ganho bruto", "Base trib.", "IR (R$)"]
    aligns = ["left", "left", "center", "right", "right", "right"]
    rows = [[""] * 6]  # header placeholder

    data_rows = []
    for b in espolio.bens:
        r = b.resultado
        ok = r.get("status") == StatusApuracao.OK
        ganho = r.get("ganho_bruto") or r.get("ganho_diferido")
        base  = r.get("base_tributavel")
        ir    = r.get("ir_espolio")
        regime_short = {
            Regime.TRANSMISSAO_HERDEIRO:         "Transmissão",
            Regime.VENDA_PELO_ESPOLIO:           "Venda",
            Regime.CESSAO_DIREITOS_HEREDITARIOS: "Cessão",
            Regime.FORA_DO_ESPOLIO:              "Fora do monte",
            Regime.INDETERMINADO:                "Indeterminado",
        }.get(b.regime, b.regime.value)
        data_rows.append((ok, [
            b.identificacao, regime_short,
            "OK" if ok else "PENDENTE",
            _m(ganho) if ganho is not None else "—",
            _m(base)  if base  is not None else "—",
            _m(ir)    if ir    is not None else "pendente",
        ]))

    tbl = doc.add_table(rows=1 + len(data_rows), cols=6)
    tbl.style = "Table Grid"
    _table_header_row(tbl.rows[0], cols)
    for i, (ok, vals) in enumerate(data_rows, 1):
        bg = _BG_TBL if i % 2 == 0 else None
        color = None if ok else _WARN
        _table_data_row(tbl.rows[i], vals, aligns=aligns, bg=bg, color=color)

    doc.add_paragraph()


def _bloco_herdeiros(doc: Document, espolio: Espolio):
    _heading(doc, "3. IR dos Herdeiros")
    tem = False

    for b in espolio.bens:
        r = b.resultado
        if r.get("status") != StatusApuracao.OK:
            continue

        if b.regime == Regime.TRANSMISSAO_HERDEIRO:
            tem = True
            _heading(doc, b.identificacao, level=2)
            _aviso(doc,
                   "IR no espólio = R$ 0,00",
                   "Herança é rendimento isento (Lei 7.713/88 Art. 6º XVI). "
                   "O ganho fica diferido ao herdeiro.", "ok")
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = "Table Grid"
            _table_header_row(tbl.rows[0], ["Campo", "Valor"])
            _table_data_row(tbl.rows[1], ["Base de custo do herdeiro",
                                           _m(r.get("base_custo_herdeiro"))])
            _table_data_row(tbl.rows[2], ["Base de custo da meação",
                                           _m(r.get("base_custo_meeiro"))],
                            bg=_BG_TBL)
            _table_data_row(tbl.rows[3], ["Ganho diferido (venda futura)",
                                           _m(r.get("ganho_diferido"))],
                            bold=True, bg=_BG_TOT, color=_NAVY)
            doc.add_paragraph()

        elif b.regime == Regime.CESSAO_DIREITOS_HEREDITARIOS:
            tem = True
            _heading(doc, b.identificacao + " — Cessão de Direitos", level=2)
            _aviso(doc,
                   "IR nas pessoas físicas dos cedentes",
                   "Cada herdeiro cedente deve recolher o GCAP na sua DIRPF. "
                   "O espólio não recolhe IR nesta operação.", "ambar")
            ir_h = r.get("ir_herdeiros") or {}
            if ir_h:
                tbl = doc.add_table(rows=1 + len(ir_h), cols=5)
                tbl.style = "Table Grid"
                _table_header_row(tbl.rows[0],
                    ["Herdeiro", "Fração", "Ganho rateado",
                     "Base tributável", "IR (R$)"])
                for i, (nome, d) in enumerate(ir_h.items(), 1):
                    bg = _BG_TBL if i % 2 == 0 else None
                    _table_data_row(tbl.rows[i], [
                        nome, f"{d['fracao']:.4f}",
                        _m(d['ganho']), _m(d['base_tributavel']), _m(d['ir']),
                    ], aligns=["left","right","right","right","right"], bg=bg)
            doc.add_paragraph()

    if not tem:
        p = doc.add_paragraph()
        _run(p, "Nenhum bem conclusivo com repercussão fiscal para os herdeiros.",
             color=_GRAY)


def _bloco_fr(doc: Document, espolio: Espolio):
    tem = any(b.resultado.get("fatores_reducao") is not None
              for b in espolio.bens)
    if not tem:
        return

    _heading(doc, "4. Memória dos Fatores de Redução")
    _aviso(doc, "Verificação obrigatória",
           "Conferir contra o programa GCAP oficial da RFB. Em divergência, "
           "prevalece o GCAP.", "info")

    for b in espolio.bens:
        fr = b.resultado.get("fatores_reducao")
        if fr is None:
            continue
        _heading(doc, b.identificacao, level=2)
        if not fr.aplicavel:
            p = doc.add_paragraph()
            _run(p, fr.motivo_inaplicavel, color=_GRAY)
            continue
        for i, passo in enumerate(fr.memoria, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            _run(p, f"Passo {i}: ", bold=True, color=_NAVY, size_pt=9)
            _run(p, passo, size_pt=9)
        # resumo numérico
        p_fr = doc.add_paragraph()
        p_fr.paragraph_format.left_indent = Cm(0.5)
        _run(p_fr,
             f"M1={fr.m1}  FR1={fr.fr1:.6f}  |  "
             f"M2={fr.m2}  FR2={fr.fr2:.6f}  |  "
             f"Red. 7.713/88={fr.reducao_7713:.4f}  |  "
             f"Base final={_m(fr.base_calculo_final)}",
             size_pt=9, color=_BLUE)
        doc.add_paragraph()


def _bloco_ia(doc: Document, analise_ia: str):
    if not analise_ia or not analise_ia.strip():
        return
    _heading(doc, "5. Análise Auditorial (IA)")
    _aviso(doc, "Nota",
           "Leitura redigida pelo Claude Sonnet sobre o relatório do motor. "
           "Não recalcula valores; comenta o que foi apurado.", "info")
    for bloco in analise_ia.split("\n\n"):
        bloco = bloco.strip()
        if not bloco:
            continue
        linhas = [l.strip() for l in bloco.split("\n") if l.strip()]
        if all(l.startswith(("-", "•", "*")) for l in linhas):
            for l in linhas:
                p = doc.add_paragraph(style="List Bullet")
                _run(p, l.lstrip("-•* ").strip())
        else:
            p = doc.add_paragraph()
            texto = " ".join(linhas)
            partes = texto.split("**")
            for i, parte in enumerate(partes):
                _run(p, parte, bold=(i % 2 == 1))


def _bloco_base_legal(doc: Document, secao: int = 5):
    _heading(doc, f"{secao}. Base Legal e Avisos")
    _aviso(doc, "Aviso de uso",
           "Este laudo é software de apoio à auditoria. Não substitui o juízo "
           "do contador e do advogado responsáveis. Valide os fatores de redução "
           "contra o GCAP oficial da RFB antes de qualquer entrega.", "ok")

    normas = [
        ("Lei 9.532/97 Art. 23",
         "partilha do espólio; opção entre valor histórico (§2º) e valor de mercado (§1º)."),
        ("Lei 11.196/05 Arts. 39-40",
         "fatores de redução FR1/FR2 para imóveis adquiridos antes de nov/2005."),
        ("Lei 7.713/88 Art. 18",
         "redução por ano de aquisição anterior a 1989."),
        ("Lei 7.713/88 Art. 6º XVI",
         "herança e doação: rendimento isento do IR."),
        ("Lei 8.981/95 Art. 21 (Lei 13.259/16)",
         "alíquotas progressivas do ganho de capital (15% a 22,5%)."),
        ("IN SRF 599/2005 · RIR/2020 (Dec. 9.580/2018) Art. 150",
         "disposições complementares."),
    ]
    for base, desc in normas:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        _run(p, base + ": ", bold=True, color=_NAVY, size_pt=9)
        _run(p, desc, color=_GRAY, size_pt=9)

    doc.add_paragraph()
    p_ass = doc.add_paragraph()
    _run(p_ass, "_" * 35 + "\n", color=_NAVY)
    _run(p_ass, "Robson Alain Veloso\n", bold=True, color=_NAVY)
    _run(p_ass, "CRC/TO-002032/O-5 T-GO  ·  ORGATEC Contabilidade e Auditoria",
         color=_GRAY, size_pt=9)


# ───────────────────────────────────────────────────────────────────────
# Geração
# ───────────────────────────────────────────────────────────────────────

def gerar_laudo_docx(
    espolio: Espolio,
    caminho_saida: str,
    analise_ia: Optional[str] = None,
) -> str:
    """
    Gera o laudo em formato Word (.docx) e salva em caminho_saida.
    Retorna o caminho do arquivo gerado.
    """
    motor = MotorApuracao(espolio)
    resumo = motor.apurar()

    doc = Document()

    # Margens A4
    for sec in doc.sections:
        sec.page_width  = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin   = sec.right_margin  = Cm(2.0)
        sec.top_margin    = sec.bottom_margin = Cm(2.0)

    _bloco_capa(doc, espolio)
    _bloco_resumo(doc, resumo, espolio)
    _bloco_apuracao(doc, espolio)
    _bloco_herdeiros(doc, espolio)
    _bloco_fr(doc, espolio)
    if analise_ia:
        _bloco_ia(doc, analise_ia)
    _bloco_base_legal(doc, secao=6 if analise_ia else 5)

    doc.save(caminho_saida)
    return caminho_saida
